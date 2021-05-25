"""
created by Brad Hillier - BradleyHillier@icloud.com
generate the morning meeting document for Sealegs Kayaking Adventures 
in Ladysmith, BC
"""



from dataclasses import dataclass, field
from datetime import datetime
from bs4 import BeautifulSoup
from selenium import webdriver
from docxtpl import DocxTemplate
from docx import Document
from time import sleep
import requests
import dateparser
import pytz


    


@dataclass
class Weather:
    date: datetime
    description: str
    temp: str
    wind: str

    def __str__(self):
        attrs = [
            self.date.strftime('%-I %p'),
            self.description,
            self.temp + u'\N{DEGREE SIGN}' + 'C',
            self.wind
        ]
        return ' - '.join(attrs)


@dataclass
class Tide:
    time: datetime
    meters: int
    feet: int

    def __str__(self):
        time = self.time.strftime('%-I:%M %p')
        return f'{self.meters}m @ {time}'


@dataclass 
class Booking:
    title: str
    start_at: datetime
    end_at: datetime
    description: str = field(default=None)

    def __str__(self):
        start = self.start_at.strftime('%-I %p')
        end = self.end_at.strftime('%-I %p')
        des = self.description.replace('\n', ' ')
        return f'{self.title} from {start} to {end} - {des}'



def main():
    datestring = datetime.now().strftime('%A - %B %d - %Y')
    tides = get_tides()
    weather = get_weather(10, 17)
    bookings = get_bookings()

    document = Document()
    document.add_heading('Daily Safety Meeting', 0)
    document.add_paragraph(datestring)

    document.add_heading('Attendants', 2)
    for section in ['Safety Topic', 'General Notes/Maintenance']:
        document.add_heading(section, 2)
        document.add_paragraph('\n' * 2)

    document.add_heading('Tides', 2)
    tides_table = document.add_table(rows=0, cols=3)
    for tide in tides:
        row_cells = tides_table.add_row().cells
        row_cells[0].text = tide.time.strftime('%-I:%M %p')
        row_cells[1].text = f'{tide.meters} meters'
        row_cells[2].text = f'{tide.feet} feet'

    document.add_heading('Weather', 2)
    wx_table = document.add_table(rows=4, cols=len(weather))
    for i in range(len(wx_table.columns)):
        col = wx_table.columns[i]
        col.cells[0].text = weather[i].date.strftime('%-I %p')
        col.cells[1].text = emoji_from_description(weather[i].description)
        col.cells[2].text = weather[i].temp + u'\N{DEGREE SIGN}' + 'C',
        col.cells[3].text = weather[i].wind

    document.add_heading('Bookings', 2)
    bookings_table = document.add_table(rows=1, cols=2)
    for booking in bookings:
        row_cells = bookings_table.add_row().cells
        start = booking.start_at.strftime('%-I %p')
        end = booking.end_at.strftime('%-I %p')
        row_cells[0].text = f'{start} to {end}'
        row_cells[1].text = booking.title


    document.save('test.docx')



def get_tides() -> list:
    '''Retreive today's tides from 'Fisheries and Oceans Canada'.'''
    now = datetime.now()
    url = 'https://tides.gc.ca/eng/station?type=0&date={}%2F{}%2F{}&sid=7460&tz=PDT&pres=1'.format(
        now.year, '%02d'%now.month, now.day)
    res = requests.get(url)
    soup = BeautifulSoup(res.content, 'html.parser')

    time = soup.table.tbody.findAll('td', class_='time')
    meters = soup.table.tbody.findAll('td', class_='heightMeters')
    feet = soup.table.tbody.findAll('td', class_='heightFeet')

    tides = [Tide(dateparser.parse(time[i].text), meters[i].text, feet[i].text,) \
            for i in range(len(time))]

    return [tide for tide in tides]
        

# No APIs I looked at offered hourly weather data for free, and content is
# loaded dynamically on "theweathernetwork.com".  because of this it was 
# necessary to use web browser automation.
def get_weather(start_time: int, end_time: int) -> list:
    '''
    Retrieve today's hourly weather data from 'The Weather Network' 
    for the provided time window

    Arguments:
    start_time -- the hour (0-23) to begin retrieving data
    end_time -- the hour (0-23) to finish grabbing data (inclusive)
    '''
    url = 'https://www.theweathernetwork.com/ca/hourly-weather-forecast/british-columbia/ladysmith'
    browser = webdriver.Safari()
    browser.get(url)
    browser.implicitly_wait(10)
    raw_hourly_wx = browser.find_elements_by_class_name('wxColumn-hourly')
    wx_legend = browser.find_elements_by_class_name('legendColumn')
    wind_index = find_wind_index(wx_legend)

    # webelements on hidden tables were showing up with empty text in Chrome. 
    # Clicking the button to show the next table solved this issue. 
    # Issue did not exist when using the Safari driver. 
    # Six hours of weather data are displayed in each table
    hourly_wx = []
    hours_per_table = 6
    num_of_tables = int(len(raw_hourly_wx) / hours_per_table)
    for i in range(num_of_tables):
        for j in range(hours_per_table):
            curr_hour_wx = parse_weather_from_webelement(
                raw_hourly_wx[hours_per_table * i + j], 
                wind_index
            )
            if curr_hour_wx.date.hour >= start_time and \
                curr_hour_wx.date.hour <= end_time:
                hourly_wx.append(curr_hour_wx)
            else:
                break
        else:
            browser.find_element_by_class_name('hourlyforecast_data_table_ls_next').click()
            sleep(0.5)
            continue
        break

    browser.quit()
    return hourly_wx


# Table rows are dynamically generated based on information relevant to current
# days conditions
def find_wind_index(wx_legend: webdriver.remote.webelement.WebElement) -> int:
    '''
    Helper Function for get_weather
    Determine which table row contains wind data 
    '''
    for i in range(len(wx_legend)):
        if wx_legend[i].text == 'Wind (km/h)':
            return i


def parse_weather_from_webelement(column: webdriver.remote.webelement.WebElement,
                                  wind_index: int) -> Weather:
    '''
    Helper function for get_weather
    Parses data from a WebElement and creates a Weather object
    '''
    weekday = column.find_element_by_class_name('day').text
    time = column.find_element_by_class_name('date').text

    date = next_occurence(weekday, time)
    description = column.find_element_by_xpath('.//*[@class="wx_description"]').text
    temp = column.find_element_by_xpath('.//*[@class="wxperiod_temp"]').text
    wind = column.find_elements_by_class_name('stripeable')[wind_index].text

    return Weather(date, description, temp, wind)


# dateparser weeks appear to start on a Saturday; when parsing without this 
# it would output the previous occurence if the weekday was less than
# now.weekday()
def next_occurence(weekday: str, time: str) -> datetime:
    '''
    Returns the next occurence of the provided weekday and time as a 
    datetime object
    '''
    now = datetime.now()
    date = dateparser.parse(' '.join([weekday, time]))
    if date.day - now.day < 0:
        return dateparser.parse(' '.join([weekday, time]), 
                                settings={'PREFER_DATES_FROM': 'future'})
    else:
        return date
    

def emoji_from_description(description: str) -> str:
    '''Return a weather emoji depicting the provided description'''

    # sun
    if description == 'Sunny': 
        return '\u2600' 

    # sun behind small cloud
    if description == 'Mainly sunny': 
        return '\U0001F324'

    # sun behind cloud
    if description in ('A mix of sun and clouds', 'Partly cloudy'): 
        return '\u26c5' 

    # sun behind large cloud
    if description in ('Cloudy with clear breaks', 'Cloudy with sunny breaks'):
        return '\U0001F325'

    # sun behind rain cloud
    if description == 'Chance of a shower':
        return '\U0001F326'

    # cloud with rain
    if description in ('Cloudy with showers', 'A few showers', 'Light rain'): 
        return '\U0001F327'

    # crescent moon
    if description in ('Clear', 'Mainly clear'): 
        return '\U0001F317' # crescent moon

    # cloud
    if description == 'Mainly cloudy': 
        return '\u2601' 

    # fog
    if description == 'Fog patches':
        return '\U0001F329'

    # red question mark
    with open('mmg.log', 'a') as f:
        now_str = datetime.now().strftime('%x %X')
        f.write(f'\n{now_str} unknown weather description "{description}"')
    return '\u2753'


def get_bookings() -> list:
    '''
    Retrieve bookings from timetree's API upcoming_events endpoint
    Returns only current days upcoming bookings
    '''
    token = '6E-18RAX47fodf5dh3TPBpylbmXr9JmZ3mBkJ0V4bqe1It0n'
    cal_id = 'Mw4DZK3sc72B'
    base_url = 'https://timetreeapis.com/'
    headers = {'accept': 'application/vnd.timetree.v1+json',
              'Authorization': f'Bearer {token}',
             }
    res = requests.get( f'{base_url}/calendars/{cal_id}/upcoming_events',
                       headers=headers)
    raw_bookings = res.json()['data']
    bookings = [create_booking_obj(raw) for raw in raw_bookings]

    return [booking for booking in bookings if \
           booking.start_at.day == datetime.now.day()]


def create_booking_obj(raw_event: dict) -> Booking:
    '''
    Helper function for get_bookings
    Converts raw booking data into Booking object
    '''
    timezone = pytz.timezone('America/Vancouver')

    title = raw_event['attributes']['title'],
    start = dateparser.parse(raw_event['attributes']['start_at']).astimezone(timezone)
    end = dateparser.parse(raw_event['attributes']['end_at']).astimezone(timezone)
    des = raw_event['attributes']['description'] or ''

    return Booking(title, start, end, des)



if __name__ == '__main__':
    main()
